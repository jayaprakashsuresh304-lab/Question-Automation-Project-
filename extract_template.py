from docx import Document
from docx.shared import Inches, Pt
import json

try:
    doc = Document(r"C:\Users\PRAJEEN\AppData\Local\Packages\5319275A.WhatsAppDesktop_cv1g1gvanyjgm\LocalState\sessions\380D00BC6FE1B5D7E953E69463CF78C28F7FFBA2\transfers\2026-12\CT1 DM.docx")
    
    print("=" * 100)
    print("COLLEGE QUESTION PAPER TEMPLATE ANALYSIS")
    print("=" * 100)
    
    # Extract all paragraphs
    print("\n\n### FULL DOCUMENT TEXT ###\n")
    para_count = 0
    for para in doc.paragraphs:
        if para.text.strip():
            para_count += 1
            alignment = para.alignment
            align_text = {None: "LEFT", 0: "LEFT", 1: "CENTER", 2: "RIGHT"}.get(alignment, "LEFT")
            
            bold = False
            font_size = "?pt"
            font_name = "?"
            
            if para.runs and len(para.runs) > 0:
                bold = para.runs[0].bold
                if para.runs[0].font.size:
                    font_size = f"{para.runs[0].font.size.pt}pt"
                if para.runs[0].font.name:
                    font_name = para.runs[0].font.name
            
            print(f"\n[P{para_count}] {align_text:8} | {font_name:12} {font_size:6} | Bold:{bold}")
            print(f"     {para.text[:100]}")
    
    # Extract tables
    print("\n\n### TABLES ###\n")
    print(f"Total tables: {len(doc.tables)}")
    
    for t_idx, table in enumerate(doc.tables):
        print(f"\nTable {t_idx + 1}: {len(table.rows)} rows x {len(table.columns)} columns")
        for row_idx, row in enumerate(table.rows[:5]):  # First 5 rows
            print(f"  Row {row_idx}:", end=" ")
            for cell in row.cells:
                print(f"[{cell.text[:40]}]", end=" ")
            print()
    
    # Save to JSON for web app
    template_data = {
        "total_paragraphs": para_count,
        "total_tables": len(doc.tables),
        "pages": 1,
    }
    
    with open("template_info.json", "w") as f:
        json.dump(template_data, f, indent=2)
    
    print("\n\nTemplate info saved to template_info.json")
    
except Exception as e:
    print(f"Error: {e}")
    import traceback
    traceback.print_exc()
