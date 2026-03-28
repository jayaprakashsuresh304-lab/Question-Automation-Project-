from __future__ import annotations

import csv
import json
from io import BytesIO
from functools import wraps
import re
import sqlite3
from zipfile import BadZipFile, ZipFile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Any
from uuid import uuid4

from flask import Flask, flash, g, jsonify, redirect, render_template, request, send_file, session, url_for
from markupsafe import Markup, escape
from werkzeug.security import check_password_hash, generate_password_hash
from werkzeug.utils import secure_filename

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

BASE_DIR = Path(__file__).resolve().parent
DB_PATH = BASE_DIR / "question_bank.db"
UPLOAD_DIR = BASE_DIR / "static" / "uploads"
UPLOAD_STAGE_DIR = BASE_DIR / "tmp_uploads"
ALLOWED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".svg"}

app = Flask(__name__)
app.config["SECRET_KEY"] = "dev-key-change-in-production"


def ensure_schema_extensions(conn: sqlite3.Connection) -> None:
    """Lightweight migration for newer question content features."""
    columns = {
        row[1]
        for row in conn.execute("PRAGMA table_info(questions)").fetchall()
    }
    if "diagram_path" not in columns:
        conn.execute("ALTER TABLE questions ADD COLUMN diagram_path TEXT")
        conn.commit()

    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            full_name TEXT NOT NULL,
            email TEXT UNIQUE NOT NULL,
            role TEXT DEFAULT 'teacher',
            is_active BOOLEAN DEFAULT 1,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
        """
    )

    admin_exists = conn.execute(
        "SELECT 1 FROM users WHERE email = ? LIMIT 1",
        ("admin@mzcollege.com",),
    ).fetchone()
    if not admin_exists:
        conn.execute(
            """
            INSERT INTO users (username, password_hash, full_name, email, role, is_active)
            VALUES (?, ?, ?, ?, ?, 1)
            """,
            (
                "admin",
                generate_password_hash("admin123"),
                "Administrator",
                "admin@mzcollege.com",
                "admin",
            ),
        )
    else:
        conn.execute(
            "UPDATE users SET role = 'admin', is_active = 1 WHERE email = ?",
            ("admin@mzcollege.com",),
        )
    conn.commit()

    paper_columns = {
        row[1]
        for row in conn.execute("PRAGMA table_info(question_papers)").fetchall()
    }
    if paper_columns:
        if "paper_data_json" not in paper_columns:
            conn.execute("ALTER TABLE question_papers ADD COLUMN paper_data_json TEXT")
        if "edited_payload_json" not in paper_columns:
            conn.execute("ALTER TABLE question_papers ADD COLUMN edited_payload_json TEXT")
        conn.commit()


def save_diagram_file(file_storage) -> str | None:
    if not file_storage or not file_storage.filename:
        return None

    original = secure_filename(file_storage.filename)
    suffix = Path(original).suffix.lower()
    if suffix not in ALLOWED_IMAGE_EXTENSIONS:
        return None

    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"{uuid4().hex}{suffix}"
    target = UPLOAD_DIR / filename
    file_storage.save(target)
    return f"uploads/{filename}"


def save_diagram_blob(blob: bytes, suffix: str) -> str | None:
    if not blob:
        return None
    normalized_suffix = (suffix or "").lower()
    if normalized_suffix not in ALLOWED_IMAGE_EXTENSIONS:
        return None

    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"{uuid4().hex}{normalized_suffix}"
    target = UPLOAD_DIR / filename
    target.write_bytes(blob)
    return f"uploads/{filename}"


def split_diagram_paths(value: str | None) -> list[str]:
    raw = (value or "").strip()
    if not raw:
        return []
    return [p.strip() for p in raw.split("|") if p and p.strip()]


def pack_diagram_paths(paths: list[str]) -> str:
    seen: set[str] = set()
    cleaned: list[str] = []
    for p in paths:
        token = (p or "").strip()
        if not token or token in seen:
            continue
        seen.add(token)
        cleaned.append(token)
    return "|".join(cleaned)


def delete_diagram_paths(value: str | None) -> None:
    for path in split_diagram_paths(value):
        delete_diagram_file(path)


app.jinja_env.globals["split_diagram_paths"] = split_diagram_paths


def delete_diagram_file(relative_path: str | None) -> None:
    if not relative_path:
        return
    target = BASE_DIR / "static" / relative_path
    if target.exists() and target.is_file():
        target.unlink()


def save_staged_upload(rows: list[dict[str, str]]) -> str:
    UPLOAD_STAGE_DIR.mkdir(parents=True, exist_ok=True)
    token = uuid4().hex
    target = UPLOAD_STAGE_DIR / f"{token}.json"
    target.write_text(json.dumps(rows, ensure_ascii=True), encoding="utf-8")
    return token


def load_staged_upload(token: str) -> list[dict[str, str]] | None:
    if not token:
        return None
    target = UPLOAD_STAGE_DIR / f"{token}.json"
    if not target.exists():
        return None
    try:
        return json.loads(target.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return None


def delete_staged_upload(token: str | None, cleanup_diagrams: bool = True) -> None:
    if not token:
        return
    target = UPLOAD_STAGE_DIR / f"{token}.json"
    if cleanup_diagrams and target.exists() and target.is_file():
        try:
            staged = json.loads(target.read_text(encoding="utf-8"))
            if isinstance(staged, list):
                for row in staged:
                    if isinstance(row, dict):
                        delete_diagram_paths(row.get("diagram_path"))
        except (json.JSONDecodeError, OSError):
            pass
    if target.exists() and target.is_file():
        target.unlink()


def normalize_column_name(name: str) -> str:
    key = name.strip().lower().replace(" ", "_").replace("-", "_")
    key = key.replace(".", "").replace(":", "")
    aliases = {
        "question": "question_text",
        "questiontext": "question_text",
        "question_text": "question_text",
        "co_mapping": "co",
        "course_outcome": "co",
        "bloomstaxonomylevel": "bt_level",
        "blooms_taxonomy_level": "bt_level",
        "blooms_level": "bt_level",
        "btlevel": "bt_level",
        "unit_no": "unit",
        "unit_number": "unit",
        "qno": "q_no",
        "q_no": "q_no",
        "question_no": "q_no",
    }
    return aliases.get(key, key)


def extract_marks(value: str) -> int | None:
    m = re.search(r"\b(2|13|15|16)\b", value)
    if not m:
        return None
    return int(m.group(1))


def detect_marks_from_heading(value: str) -> int | None:
    text = value.strip().lower()
    if not text:
        return None

    part_a = re.search(r"\bpart\s*[-:]?\s*a\b", text)
    part_b = re.search(r"\bpart\s*[-:]?\s*b\b", text)
    part_c = re.search(r"\bpart\s*[-:]?\s*c\b", text)

    if part_a or re.search(r"\b2\s*marks?\b", text):
        return 2
    if part_b or re.search(r"\b13\s*marks?\b", text):
        return 13
    if part_c:
        # Keep compatibility with existing paper format; allow explicit 16 if heading says so.
        if re.search(r"\b16\s*marks?\b", text):
            return 16
        return 15
    if re.search(r"\b15\s*marks?\b", text):
        return 15
    if re.search(r"\b16\s*marks?\b", text):
        return 16
    return None


def normalize_question_text(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", (text or "").strip())
    # Remove leading numbering tokens like 1., 1), Q1), 11.a), etc.
    cleaned = re.sub(r"^(?:q\s*)?\d+[a-z]?\s*[.)-]\s*", "", cleaned, flags=re.IGNORECASE)
    # Remove trailing marks annotation like (13 marks) or 13 marks.
    cleaned = re.sub(r"\(?\b(2|13|15|16)\s*marks?\b\)?\.?$", "", cleaned, flags=re.IGNORECASE).strip()
    # Remove common OCR/Word alt text artifacts when image placeholders are extracted as text.
    cleaned = re.sub(r"\bquestion\s+diagram\b", "", cleaned, flags=re.IGNORECASE).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned


def normalize_question_text_preserve_structure(text: str) -> str:
    raw = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    # Handle HTML <br> tags stored in database
    raw = raw.replace("<br>", "\n").replace("<br/>", "\n").replace("<BR>", "\n")
    lines = [re.sub(r"\s+", " ", line).strip() for line in raw.split("\n")]
    lines = [line for line in lines if line]

    return "\n".join(lines)


def int_to_roman(value: int) -> str:
    if value <= 0:
        return "i"
    mapping = [
        (1000, "m"),
        (900, "cm"),
        (500, "d"),
        (400, "cd"),
        (100, "c"),
        (90, "xc"),
        (50, "l"),
        (40, "xl"),
        (10, "x"),
        (9, "ix"),
        (5, "v"),
        (4, "iv"),
        (1, "i"),
    ]
    out = ""
    n = value
    for arabic, roman in mapping:
        while n >= arabic:
            out += roman
            n -= arabic
    return out or "i"


def build_marked_segments(text: str) -> list[tuple[str, str]]:
    compact = re.sub(r"\s+", " ", (text or "")).strip()
    if not compact:
        return []

    marks_matches = list(re.finditer(r"\(\s*(\d{1,2})\s*marks?\s*\)", compact, flags=re.IGNORECASE))
    if len(marks_matches) < 2:
        return []

    chunks: list[tuple[str, str]] = []
    start = 0
    for match in marks_matches:
        segment_text = compact[start:match.end()].strip(" .;,-")
        start = match.end()
        if not segment_text:
            continue

        segment_text = re.sub(r"^\d+\s*[.)]\s*", "", segment_text).strip()
        chunks.append((segment_text, match.group(1)))

    normalized: list[tuple[str, str]] = []
    for idx, (segment_text, marks) in enumerate(chunks, start=1):
        if not re.match(r"^\s*\(?([ivxlcdm]+|[a-z])\)?\s*[.)]\s+", segment_text, flags=re.IGNORECASE):
            segment_text = f"{int_to_roman(idx)}) {segment_text}"
        normalized.append((segment_text, marks))

    return normalized


def infer_total_marks_from_text(text: str, fallback_marks: int | None = None) -> int | None:
    raw = text or ""

    # Use strict explicit-mark patterns only, to avoid false positives from values like 2-2'.
    strict_match = re.search(r"\(?\b(2|13|15|16)\b\s*marks?\)?", raw, flags=re.IGNORECASE)
    if strict_match:
        return int(strict_match.group(1))

    trailing_match = re.search(r"\b(2|13|15|16)\b\s*$", raw.strip())
    if trailing_match:
        return int(trailing_match.group(1))

    marked_segments = build_marked_segments(raw)
    if marked_segments:
        total = sum(int(marks) for _, marks in marked_segments)
        if total in {2, 13, 15, 16}:
            return total

    if fallback_marks in {2, 13, 15, 16}:
        return fallback_marks
    return None


def prepare_question_text_for_display(text: str | None) -> str:
    cleaned = normalize_question_text_preserve_structure(text or "")
    if not cleaned:
        return ""

    marked_segments = build_marked_segments(cleaned)
    if marked_segments:
        segment_lines: list[str] = []
        for segment, _ in marked_segments:
            pretty = re.sub(r"\s+(?=(?:\(?[ivxlcdm]+\)?|[a-z])\s*[.)]\s+)", "\n", segment, flags=re.IGNORECASE)
            pretty = re.sub(r"\s+(?=\d+\s*[.)]\s+)", "\n", pretty)
            segment_lines.append(pretty.strip())
        return "\n".join(segment_lines)

    # Backward compatibility for older rows where split sub-questions were flattened.
    if "\n" not in cleaned:
        cleaned = re.sub(r"\s+(?=(?:\(?[ivxlcdm]+\)?|[a-z])\s*[.)]\s+)", "\n", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\s+(?=\([ivxlcdm]+\)\s+)", "\n", cleaned, flags=re.IGNORECASE)

        # Also split legacy numeric subparts like: "... (8 marks) 2. ... (5 marks)"
        cleaned = re.sub(
            r"(\(?\b\d{1,2}\s*(?:marks?)\b\)?\.?)(\s+)(?=\d+\s*[.)]\s+)",
            r"\1\n",
            cleaned,
            flags=re.IGNORECASE,
        )
        cleaned = re.sub(r"\s+(?=\d+\s*[.)]\s+)", "\n", cleaned)

    lines = [line.strip() for line in cleaned.split("\n") if line.strip()]
    if len(lines) > 1:
        converted_lines: list[str] = []
        for line in lines:
            match = re.match(r"^\s*(\d+)\s*[.)]\s+(.+)$", line)
            if match:
                converted_lines.append(f"{int_to_roman(int(match.group(1)))}) {match.group(2).strip()}")
            else:
                converted_lines.append(line)
        cleaned = "\n".join(converted_lines)

    return cleaned


def format_question_text_for_display(text: str | None) -> Markup:
    cleaned = prepare_question_text_for_display(text)
    if not cleaned:
        return Markup("")

    return Markup(escape(cleaned).replace("\n", "<br>"))


def format_split_marks_for_display(text: str | None, fallback_marks: int | str = "") -> Markup:
    cleaned = prepare_question_text_for_display(text)
    if not cleaned:
        return Markup(escape(str(fallback_marks)))

    marked_segments = build_marked_segments(cleaned)
    if marked_segments:
        return Markup("<br>".join(escape(marks) for _, marks in marked_segments))

    lines = [line.strip() for line in cleaned.split("\n") if line.strip()]
    split_marks: list[str] = []

    for line in lines:
        subq_match = re.match(r"^\s*\(?([ivxlcdm]+|[a-z]|\d+)\)?\s*[.)]\s+(.+)$", line, flags=re.IGNORECASE)
        content = subq_match.group(2).strip() if subq_match else line

        marks_match = re.search(r"\(?\s*(\d{1,2})\s*(?:marks?)?\s*\)?\s*[.)]?$", content, flags=re.IGNORECASE)
        if marks_match:
            split_marks.append(marks_match.group(1))

    if split_marks:
        return Markup("<br>".join(escape(m) for m in split_marks))
    return Markup(escape(str(fallback_marks)))


app.jinja_env.filters["format_question_text"] = format_question_text_for_display
app.jinja_env.filters["format_split_marks"] = format_split_marks_for_display


def split_marks_from_text(text: str | None, fallback_marks: int) -> list[str]:
    marked_segments = build_marked_segments(text or "")
    if marked_segments:
        return [marks for _, marks in marked_segments]

    cleaned = prepare_question_text_for_display(text)
    lines = [line.strip() for line in cleaned.split("\n") if line.strip()]
    values: list[str] = []
    for line in lines:
        subq_match = re.match(r"^\s*\(?([ivxlcdm]+|[a-z]|\d+)\)?\s*[.)]\s+(.+)$", line, flags=re.IGNORECASE)
        content = subq_match.group(2).strip() if subq_match else line
        marks_match = re.search(r"\(?\s*(\d{1,2})\s*(?:marks?)?\s*\)?\s*[.)]?$", content, flags=re.IGNORECASE)
        if marks_match:
            values.append(marks_match.group(1))

    return values or [str(fallback_marks)]


def serialize_row(row: sqlite3.Row | dict[str, Any] | None) -> dict[str, Any]:
    if row is None:
        return {}
    if isinstance(row, sqlite3.Row):
        return dict(row)
    return dict(row)


def save_generated_paper(
    bank: sqlite3.Row,
    college_name: str,
    department: str,
    semester: str,
    exam_type: str,
    exam_date: str,
    exam_time: str,
    part_a_count: int,
    selected_part_a: list[sqlite3.Row],
    selected_part_b: list[sqlite3.Row],
    selected_part_c: list[sqlite3.Row],
) -> int:
    db = get_db()
    created_by = int(session.get("user_id") or 1)

    paper_title = f"{exam_type or 'Question Paper'} - {exam_date}".strip(" -")
    paper_data = {
        "college_name": college_name,
        "department": department,
        "semester": semester,
        "exam_type": exam_type,
        "exam_date": exam_date,
        "exam_time": exam_time,
        "part_a_count": part_a_count,
    }

    db.execute(
        """
        INSERT INTO question_papers
        (subject_id, semester_id, paper_title, exam_type, exam_date, exam_time, created_by, status, total_marks, paper_data_json)
        VALUES (?, ?, ?, ?, ?, ?, ?, 'draft', 100, ?)
        """,
        (
            bank["subject_id"],
            bank["semester_id"],
            paper_title,
            exam_type,
            exam_date,
            exam_time,
            created_by,
            json.dumps(paper_data, ensure_ascii=True),
        ),
    )
    paper_id = int(db.execute("SELECT last_insert_rowid()").fetchone()[0])

    for idx, q in enumerate(selected_part_a, start=1):
        db.execute(
            """
            INSERT INTO paper_questions (paper_id, question_id, part, question_number, has_choice, choice_question_id)
            VALUES (?, ?, 'A', ?, 0, NULL)
            """,
            (paper_id, q["id"], idx),
        )

    for i in range(5):
        qa = selected_part_b[i]
        qb = selected_part_b[i + 5] if (i + 5) < len(selected_part_b) else qa
        db.execute(
            """
            INSERT INTO paper_questions (paper_id, question_id, part, question_number, has_choice, choice_question_id)
            VALUES (?, ?, 'B', ?, 1, ?)
            """,
            (paper_id, qa["id"], 11 + i, qb["id"]),
        )

    qa_c = selected_part_c[0]
    qb_c = selected_part_c[1] if len(selected_part_c) > 1 else selected_part_c[0]
    db.execute(
        """
        INSERT INTO paper_questions (paper_id, question_id, part, question_number, has_choice, choice_question_id)
        VALUES (?, ?, 'C', 16, 1, ?)
        """,
        (paper_id, qa_c["id"], qb_c["id"]),
    )

    db.commit()
    return paper_id


def get_paper_render_context(paper_id: int) -> dict[str, Any] | None:
    db = get_db()
    paper = db.execute(
        """
        SELECT qp.*, s.subject_code, s.subject_name, sem.semester_number, sem.academic_year,
               qb.id AS question_bank_id
        FROM question_papers qp
        JOIN subjects s ON qp.subject_id = s.id
        JOIN semesters sem ON qp.semester_id = sem.id
        LEFT JOIN question_banks qb ON qb.subject_id = qp.subject_id AND qb.semester_id = qp.semester_id
        WHERE qp.id = ?
        """,
        (paper_id,),
    ).fetchone()
    if not paper:
        return None

    raw_meta = paper["paper_data_json"] or "{}"
    try:
        metadata = json.loads(raw_meta)
        if not isinstance(metadata, dict):
            metadata = {}
    except json.JSONDecodeError:
        metadata = {}

    raw_saved = paper["edited_payload_json"] or ""
    try:
        saved_payload = json.loads(raw_saved) if raw_saved else []
        if not isinstance(saved_payload, list):
            saved_payload = []
    except json.JSONDecodeError:
        saved_payload = []

    rows = db.execute(
        """
        SELECT pq.*, 
               q.id AS q_id, q.question_text AS q_question_text, q.co_mapping AS q_co_mapping,
               q.bt_level AS q_bt_level, q.difficulty AS q_difficulty, q.diagram_path AS q_diagram_path,
               cq.id AS cq_id, cq.question_text AS cq_question_text, cq.co_mapping AS cq_co_mapping,
               cq.bt_level AS cq_bt_level, cq.difficulty AS cq_difficulty, cq.diagram_path AS cq_diagram_path
        FROM paper_questions pq
        JOIN questions q ON pq.question_id = q.id
        LEFT JOIN questions cq ON pq.choice_question_id = cq.id
        WHERE pq.paper_id = ?
        ORDER BY pq.part ASC, pq.question_number ASC, pq.id ASC
        """,
        (paper_id,),
    ).fetchall()

    part_a_questions: list[dict[str, Any]] = []
    part_b_questions: list[dict[str, Any]] = []
    part_c_questions: list[dict[str, Any]] = []

    for row in rows:
        q_primary = {
            "id": row["q_id"],
            "question_text": row["q_question_text"],
            "co_mapping": row["q_co_mapping"],
            "bt_level": row["q_bt_level"],
            "difficulty": row["q_difficulty"],
            "diagram_path": row["q_diagram_path"],
        }
        q_choice = {
            "id": row["cq_id"],
            "question_text": row["cq_question_text"],
            "co_mapping": row["cq_co_mapping"],
            "bt_level": row["cq_bt_level"],
            "difficulty": row["cq_difficulty"],
            "diagram_path": row["cq_diagram_path"],
        }

        if row["part"] == "A":
            part_a_questions.append(q_primary)
        elif row["part"] == "B":
            part_b_questions.append(q_primary)
            if row["has_choice"] and q_choice["id"] is not None:
                part_b_questions.append(q_choice)
        elif row["part"] == "C":
            part_c_questions.append(q_primary)
            if row["has_choice"] and q_choice["id"] is not None:
                part_c_questions.append(q_choice)

    bank = {
        "id": paper["question_bank_id"],
        "subject_id": paper["subject_id"],
        "semester_id": paper["semester_id"],
        "subject_code": paper["subject_code"],
        "subject_name": paper["subject_name"],
        "semester_number": paper["semester_number"],
        "academic_year": paper["academic_year"],
    }

    return {
        "paper_id": paper_id,
        "bank": bank,
        "college_name": metadata.get("college_name", ""),
        "department": metadata.get("department", ""),
        "semester": metadata.get("semester", f"Semester {paper['semester_number']}"),
        "exam_type": metadata.get("exam_type", paper["exam_type"] or ""),
        "exam_date": metadata.get("exam_date", paper["exam_date"] or ""),
        "exam_time": metadata.get("exam_time", paper["exam_time"] or ""),
        "part_a_questions": part_a_questions,
        "part_b_questions": part_b_questions,
        "part_c_questions": part_c_questions,
        "saved_edits_payload": saved_payload,
    }


def build_paper_docx_bytes(context: dict[str, Any]) -> BytesIO:
    if Document is None:
        raise RuntimeError("python-docx is not installed")

    doc = Document()
    doc.add_heading(context.get("exam_type") or "Question Paper", level=1)
    doc.add_paragraph(context.get("college_name") or "")
    doc.add_paragraph(
        f"Course: {context['bank']['subject_code']} - {context['bank']['subject_name']} | "
        f"Date: {context.get('exam_date') or '-'} | Time: {context.get('exam_time') or '-'}"
    )
    doc.add_paragraph(f"Department/Semester: {context.get('department') or '-'} / {context.get('semester') or '-'}")

    doc.add_heading("Part A", level=2)
    for idx, q in enumerate(context.get("part_a_questions", []), start=1):
        text = prepare_question_text_for_display(q.get("question_text") or "")
        doc.add_paragraph(f"{idx}. {text}")

    doc.add_heading("Part B", level=2)
    part_b = context.get("part_b_questions", [])
    for i in range(5):
        qa = part_b[i] if i < len(part_b) else None
        qb = part_b[i + 5] if (i + 5) < len(part_b) else qa
        if qa:
            marks = ", ".join(split_marks_from_text(qa.get("question_text"), 13))
            doc.add_paragraph(f"{11 + i}.a {prepare_question_text_for_display(qa.get('question_text'))} [{marks}]")
        if qb:
            marks = ", ".join(split_marks_from_text(qb.get("question_text"), 13))
            doc.add_paragraph(f"{11 + i}.b {prepare_question_text_for_display(qb.get('question_text'))} [{marks}]")

    doc.add_heading("Part C", level=2)
    part_c = context.get("part_c_questions", [])
    if part_c:
        doc.add_paragraph(f"16.a {prepare_question_text_for_display(part_c[0].get('question_text'))}")
    if len(part_c) > 1:
        doc.add_paragraph(f"16.b {prepare_question_text_for_display(part_c[1].get('question_text'))}")

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def is_non_question_text(text: str) -> bool:
    normalized = normalize_question_text(text)
    if len(normalized) < 8:
        return True

    lower = normalized.lower()
    collapsed = re.sub(r"[^a-z0-9]+", " ", lower).strip()

    if re.fullmatch(r"(?:or|and|part|section)\b.*", lower) and "?" not in lower:
        return True

    heading_starts = (
        "part ",
        "section ",
        "unit ",
        "course code",
        "course name",
        "department",
        "semester",
        "time",
        "max marks",
        "spr no",
    )
    if lower.startswith(heading_starts):
        return True

    heading_phrases = (
        "answer all",
        "answer any",
        "answer either",
        "choose any",
        "internal assessment",
        "question paper",
        "instructions",
        "prepared by",
        "checked by",
        "approved by",
        "affiliated to",
        "accredited by",
        "autonomous institution",
    )
    if any(phrase in lower for phrase in heading_phrases):
        return True

    header_tokens = {
        "q no",
        "question",
        "question text",
        "co",
        "co mapping",
        "bt",
        "bt level",
        "marks",
        "difficulty",
        "status",
        "choice group",
    }
    if collapsed in header_tokens:
        return True

    return False


def roman_to_int(value: str) -> int | None:
    token = (value or "").strip().upper()
    if not token:
        return None
    values = {"I": 1, "V": 5, "X": 10, "L": 50, "C": 100, "D": 500, "M": 1000}
    total = 0
    prev = 0
    for ch in reversed(token):
        current = values.get(ch)
        if current is None:
            return None
        if current < prev:
            total -= current
        else:
            total += current
            prev = current
    return total if total > 0 else None


def extract_unit_number(value: str) -> str:
    text = (value or "").strip()
    if not text:
        return ""
    match = re.search(r"\bunit\s*[-:]?\s*([ivxlcdm]+|\d+)\b", text, flags=re.IGNORECASE)
    if not match:
        return ""

    token = match.group(1)
    if token.isdigit():
        return token

    roman = roman_to_int(token)
    if roman is None:
        return ""
    return str(roman)


def detect_part_label(value: str) -> str | None:
    text = (value or "").strip().lower()
    if not text:
        return None

    part_match = re.search(r"\bpart\s*[-:]?\s*([abc])\b", text)
    if part_match:
        return part_match.group(1).upper()

    section_match = re.search(r"\bsection\s*[-:]?\s*([abc])\b", text)
    if section_match:
        return section_match.group(1).upper()

    return None


def marks_from_part(part_label: str | None) -> int | None:
    if not part_label:
        return None
    mapping = {"A": 2, "B": 13, "C": 15}
    return mapping.get(part_label.upper())


def normalize_bt_level(value: str) -> str:
    text = (value or "").strip()
    if not text:
        return ""
    match = re.search(r"\b([kKlL])\s*([1-6])\b", text)
    if not match:
        return text.upper().replace(" ", "")
    return f"{match.group(1).upper()}{match.group(2)}"


def split_level_from_text(value: str) -> tuple[str, str]:
    text = (value or "").strip()
    if not text:
        return "", ""

    match = re.search(r"\b([kKlL])\s*([1-6])\b", text)
    if not match:
        return text, ""

    level = f"{match.group(1).upper()}{match.group(2)}"
    cleaned = (text[: match.start()] + " " + text[match.end():]).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned, level


def deduplicate_parsed_rows(rows: list[dict[str, str]]) -> list[dict[str, str]]:
    deduped: list[dict[str, str]] = []
    seen: set[tuple[str, str, str]] = set()

    for row in rows:
        marks = str((row.get("marks") or "")).strip()
        unit = str((row.get("unit") or "")).strip()
        question_text = normalize_question_text(row.get("question_text") or "")
        if not question_text or is_non_question_text(question_text):
            continue

        key = (question_text.lower(), marks, unit)
        if key in seen:
            continue
        seen.add(key)

        row["question_text"] = question_text
        row["unit"] = unit
        row["marks"] = marks
        row["bt_level"] = normalize_bt_level(row.get("bt_level") or "")
        deduped.append(row)

    return deduped


def parse_rows_from_lines_state_machine(lines: list[str]) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    current_unit = ""
    current_part = ""
    current_marks: int | None = None
    current_question: dict[str, Any] | None = None

    def finalize_current() -> None:
        nonlocal current_question
        if not current_question:
            return

        fragments = current_question.get("fragments") or []
        question_text = " ".join(fragments).strip()

        sub_questions = current_question.get("sub_questions") or []
        if sub_questions:
            if question_text:
                question_text = f"{question_text}\n" + "\n".join(sub_questions)
            else:
                question_text = "\n".join(sub_questions)

        question_text, inline_level = split_level_from_text(question_text)
        bt_level = normalize_bt_level(current_question.get("bt_level") or inline_level)
        question_text = normalize_question_text_preserve_structure(question_text)

        marks_value = current_question.get("marks")
        marks = str(marks_value) if marks_value is not None else ""

        if question_text and not is_non_question_text(question_text):
            rows.append(
                {
                    "unit": current_question.get("unit") or "",
                    "marks": marks,
                    "question_text": question_text,
                    "co": "",
                    "bt_level": bt_level,
                    "difficulty": "",
                    "status": "",
                    "choice_group": "",
                }
            )

        current_question = None

    for raw in lines:
        line = re.sub(r"\s+", " ", (raw or "").strip())
        if not line:
            continue

        unit_no = extract_unit_number(line)
        if unit_no:
            finalize_current()
            current_unit = unit_no
            continue

        part_label = detect_part_label(line)
        if part_label:
            finalize_current()
            current_part = part_label
            heading_marks = detect_marks_from_heading(line)
            current_marks = heading_marks if heading_marks is not None else marks_from_part(part_label)
            continue

        heading_marks = detect_marks_from_heading(line)
        if heading_marks is not None and is_non_question_text(line):
            finalize_current()
            current_marks = heading_marks
            continue

        number_match = re.match(r"^\s*(\d{1,2})(?:\s*([a-z]))?\s*[.)]\s*(.*)$", line, flags=re.IGNORECASE)
        if number_match:
            finalize_current()
            q_no = number_match.group(1)
            suffix = number_match.group(2) or ""
            head = (number_match.group(3) or "").strip()
            if suffix:
                q_no = f"{q_no}{suffix.lower()}"
            cleaned_head, inline_level = split_level_from_text(head)
            current_question = {
                "unit": current_unit,
                "part": current_part,
                "question_no": q_no,
                "fragments": [cleaned_head] if cleaned_head else [],
                "sub_questions": [],
                "marks": current_marks,
                "bt_level": inline_level,
            }
            continue

        level_only = re.match(r"^\s*([kKlL])\s*([1-6])\s*$", line)
        if level_only and current_question:
            current_question["bt_level"] = f"{level_only.group(1).upper()}{level_only.group(2)}"
            continue

        subq_match = re.match(r"^\s*\(?([ivxlcdm]+|[a-z])\)?\s*[.)]\s+(.+)$", line, flags=re.IGNORECASE)
        if subq_match and current_question:
            tag = subq_match.group(1)
            sub_text, inline_level = split_level_from_text(subq_match.group(2).strip())
            if inline_level and not current_question.get("bt_level"):
                current_question["bt_level"] = inline_level
            current_question["sub_questions"].append(f"{tag}. {sub_text}".strip())
            continue

        if current_question:
            cleaned_line, inline_level = split_level_from_text(line)
            if inline_level and not current_question.get("bt_level"):
                current_question["bt_level"] = inline_level
            if cleaned_line and not is_non_question_text(cleaned_line):
                current_question["fragments"].append(cleaned_line)
            continue

        # Last-chance parse for loose lines.
        parsed = parse_structured_line(line)
        if parsed is None:
            parsed = parse_free_text_line(line, fallback_marks=current_marks)
        if parsed:
            rows.append(parsed)

    finalize_current()
    return deduplicate_parsed_rows(rows)


def parse_free_text_line(line: str, fallback_marks: int | None = None) -> dict[str, str] | None:
    raw = line.strip()
    if not raw:
        return None

    # Remove leading question number like: 1. / 1) / 1-
    raw = re.sub(r"^\s*\d+\s*[.)-]\s*", "", raw)

    inferred_marks = infer_total_marks_from_text(raw, fallback_marks)
    if inferred_marks is None:
        return None

    marks = str(inferred_marks)
    content = raw.strip(" -|:\t")
    content = re.sub(r"\(?\b(2|13|15|16)\s*marks?\b\)?\.?$", "", content, flags=re.IGNORECASE).strip()

    content = normalize_question_text(content)
    lower_content = content.lower()
    if lower_content.startswith("part "):
        return None
    if "answer all" in lower_content or "answer either" in lower_content:
        return None
    if not content or is_non_question_text(content):
        return None

    co = ""
    bt_level = ""

    co_match = re.search(r"\bCO\s*\d+\b", content, flags=re.IGNORECASE)
    if co_match:
        co = co_match.group(0).upper().replace(" ", "")
        content = (content[: co_match.start()] + " " + content[co_match.end():]).strip()

    bt_match = re.search(r"\bL\s*\d+\b", content, flags=re.IGNORECASE)
    if bt_match:
        bt_level = bt_match.group(0).upper().replace(" ", "")
        content = (content[: bt_match.start()] + " " + content[bt_match.end():]).strip()

    unit = ""
    unit_match = re.search(r"\bUnit\s*(\d+)\b", content, flags=re.IGNORECASE)
    if unit_match:
        unit = unit_match.group(1)
        content = (content[: unit_match.start()] + " " + content[unit_match.end():]).strip()

    return {
        "unit": unit,
        "marks": marks,
        "question_text": content,
        "co": co,
        "bt_level": bt_level,
        "difficulty": "",
        "status": "",
        "choice_group": "",
    }


def parse_docx_table_row(cells: list[str], fallback_marks: int | None = None) -> dict[str, str] | None:
    values = [c.strip() for c in cells if c and c.strip()]
    if not values:
        return None

    # Likely structure: Q.No | Question | CO | BT Level | Marks
    marks = None
    for val in reversed(values):
        marks = extract_marks(val)
        if marks is not None:
            break
    if marks is None:
        marks = infer_total_marks_from_text(" ".join(values), fallback_marks)
    if marks is None:
        return None

    if len(values) >= 2 and re.fullmatch(r"\d+", values[0]):
        values = values[1:]
    if not values:
        return None

    question_text = normalize_question_text(values[0])
    co = ""
    bt_level = ""

    for token in values[1:]:
        cleaned = token.strip()
        if not cleaned:
            continue
        if extract_marks(cleaned) is not None:
            continue
        if not bt_level and re.search(r"\b([kKlL])\s*[1-6]\b", cleaned):
            bt_level = normalize_bt_level(cleaned)
            continue
        if not co and re.search(r"\bCO\s*\d+\b", cleaned, flags=re.IGNORECASE):
            co_match = re.search(r"\bCO\s*\d+\b", cleaned, flags=re.IGNORECASE)
            if co_match:
                co = co_match.group(0).upper().replace(" ", "")
            continue

    # Legacy fallback for loosely structured columns.
    if not co and len(values) > 1:
        if re.search(r"\bCO\s*\d+\b", values[1], flags=re.IGNORECASE):
            co_match = re.search(r"\bCO\s*\d+\b", values[1], flags=re.IGNORECASE)
            if co_match:
                co = co_match.group(0).upper().replace(" ", "")
    if not bt_level and len(values) > 2:
        bt_level = normalize_bt_level(values[2])

    if not question_text or is_non_question_text(question_text):
        return None

    return {
        "unit": "",
        "marks": str(marks),
        "question_text": question_text,
        "co": co,
        "bt_level": bt_level,
        "difficulty": "",
        "status": "",
        "choice_group": "",
    }


def build_row_from_mapped_values(mapping: dict[str, int], values: list[str]) -> dict[str, str]:
    result: dict[str, str] = {}
    for field, idx in mapping.items():
        result[field] = values[idx].strip() if idx < len(values) else ""
    return result


def parse_structured_line(line: str) -> dict[str, str] | None:
    # Expected format: unit|marks|question_text|co|bt_level|difficulty|status|choice_group
    if "|" not in line:
        return None

    parts = [p.strip() for p in line.split("|")]
    if len(parts) < 3:
        return None

    while len(parts) < 8:
        parts.append("")

    question_text = normalize_question_text(parts[2])
    if not question_text or is_non_question_text(question_text):
        return None

    return {
        "unit": parts[0],
        "marks": parts[1],
        "question_text": question_text,
        "co": parts[3],
        "bt_level": parts[4],
        "difficulty": parts[5],
        "status": parts[6],
        "choice_group": parts[7],
    }


def _extract_ooxml_text(node: ET.Element, ns: dict[str, str]) -> str:
    chunks: list[str] = []
    for t in node.findall(".//w:t", ns):
        if t.text:
            chunks.append(t.text)
    return " ".join(chunks).strip()


def image_suffix_from_content_type(content_type: str | None) -> str | None:
    mapping = {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/jpg": ".jpg",
        "image/webp": ".webp",
        "image/svg+xml": ".svg",
    }
    if not content_type:
        return None
    return mapping.get(content_type.lower())


def extract_cell_diagram_paths(cell: Any, doc: Any, image_cache: dict[str, str]) -> list[str]:
    tc_xml = cell._tc.xml
    rel_ids = re.findall(r'embed="(rId\d+)"', tc_xml)
    paths: list[str] = []
    for rel_id in rel_ids:
        if rel_id in image_cache:
            paths.append(image_cache[rel_id])
            continue

        image_part = doc.part.related_parts.get(rel_id)
        if image_part is None:
            continue

        suffix = image_suffix_from_content_type(getattr(image_part, "content_type", None))
        if suffix is None:
            continue

        diagram_path = save_diagram_blob(getattr(image_part, "blob", b""), suffix)
        if diagram_path:
            image_cache[rel_id] = diagram_path
            paths.append(diagram_path)

    return split_diagram_paths(pack_diagram_paths(paths))


def extract_docx_archive_images(file_bytes: bytes) -> dict[str, str]:
    rid_to_path: dict[str, str] = {}
    try:
        with ZipFile(BytesIO(file_bytes)) as archive:
            rels_path = "word/_rels/document.xml.rels"
            if rels_path not in archive.namelist():
                return rid_to_path

            rels_xml = archive.read(rels_path)
            rels_root = ET.fromstring(rels_xml)
            rels_ns = {"r": "http://schemas.openxmlformats.org/package/2006/relationships"}

            for rel in rels_root.findall(".//r:Relationship", rels_ns):
                rel_id = rel.attrib.get("Id", "")
                target = rel.attrib.get("Target", "")
                rel_type = rel.attrib.get("Type", "")
                if not rel_id or "image" not in rel_type.lower() or not target:
                    continue

                media_path = target.replace("\\", "/")
                if not media_path.startswith("word/"):
                    media_path = f"word/{media_path.lstrip('./')}"
                if media_path not in archive.namelist():
                    continue

                suffix = Path(media_path).suffix.lower()
                blob = archive.read(media_path)
                saved = save_diagram_blob(blob, suffix)
                if saved:
                    rid_to_path[rel_id] = saved
    except Exception:
        return {}

    return rid_to_path


def parse_rows_from_docx_xml_fallback(file_bytes: bytes) -> tuple[list[dict[str, str]], str | None]:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    rows: list[dict[str, str]] = []

    try:
        with ZipFile(BytesIO(file_bytes)) as archive:
            if "word/document.xml" not in archive.namelist():
                return [], "Invalid Word file: missing word/document.xml"
            xml_bytes = archive.read("word/document.xml")
    except BadZipFile:
        return [], "Invalid DOCX file: archive is corrupted or not a real .docx file."
    except Exception as exc:
        return [], f"Unable to read DOCX archive: {exc}"

    try:
        root = ET.fromstring(xml_bytes)
    except ET.ParseError:
        return [], "Unable to parse Word XML content. Please re-save the file as .docx and upload again."

    current_marks: int | None = None
    rid_to_path = extract_docx_archive_images(file_bytes)

    # Parse table rows first (common for question banks).
    for tr in root.findall(".//w:tbl//w:tr", ns):
        cells: list[str] = []
        for tc in tr.findall("./w:tc", ns):
            text = _extract_ooxml_text(tc, ns)
            if text:
                cells.append(text)

        tr_xml = ET.tostring(tr, encoding="unicode")
        rel_ids = re.findall(r'embed="(rId\d+)"', tr_xml)
        diagram_paths: list[str] = []
        for rel_id in rel_ids:
            path = rid_to_path.get(rel_id, "")
            if path:
                diagram_paths.append(path)
        packed_paths = pack_diagram_paths(diagram_paths)

        if not cells:
            continue

        row_text = " ".join(cells)
        detected = detect_marks_from_heading(row_text)
        if detected is not None:
            current_marks = detected
            continue

        parsed = parse_docx_table_row(cells, fallback_marks=current_marks)
        if parsed:
            parsed["diagram_path"] = packed_paths
            rows.append(parsed)
        elif packed_paths and rows:
            merged = pack_diagram_paths(split_diagram_paths(rows[-1].get("diagram_path")) + split_diagram_paths(packed_paths))
            rows[-1]["diagram_path"] = merged

    # Parse paragraph lines from document body with context-aware state machine.
    paragraph_lines: list[str] = []
    for p in root.findall(".//w:body/w:p", ns):
        line = _extract_ooxml_text(p, ns)
        if line:
            paragraph_lines.append(line)

    if paragraph_lines:
        rows.extend(parse_rows_from_lines_state_machine(paragraph_lines))

    return deduplicate_parsed_rows(rows), None


def parse_rows_from_docx(file_bytes: bytes) -> tuple[list[dict[str, str]], str | None]:
    if Document is None:
        return [], "Word upload requires python-docx. Install with: pip install python-docx"

    try:
        doc = Document(BytesIO(file_bytes))
    except (KeyError, BadZipFile, ValueError) as exc:
        fallback_rows, fallback_error = parse_rows_from_docx_xml_fallback(file_bytes)
        if fallback_rows:
            return fallback_rows, None
        return [], (
            "Could not open this Word file with python-docx. "
            f"Reason: {exc}. Try re-saving as .docx (Word Document) and upload again. "
            + (fallback_error or "")
        ).strip()
    except Exception as exc:
        fallback_rows, fallback_error = parse_rows_from_docx_xml_fallback(file_bytes)
        if fallback_rows:
            return fallback_rows, None
        return [], (
            "Could not read the uploaded Word file. "
            f"Reason: {exc}. Please re-save as .docx and retry. "
            + (fallback_error or "")
        ).strip()

    rows: list[dict[str, str]] = []
    image_cache: dict[str, str] = {}

    # Parse tables and infer marks from heading rows when needed.
    for table in doc.tables:
        if not table.rows:
            continue

        first_row = [normalize_column_name(cell.text) for cell in table.rows[0].cells]
        first_row_text = " ".join(cell.text.strip() for cell in table.rows[0].cells if cell.text.strip())
        table_marks = detect_marks_from_heading(first_row_text)
        current_table_marks = table_marks
        header_fields = {
            "unit",
            "marks",
            "question_text",
            "question",
            "co",
            "bt_level",
            "difficulty",
            "status",
            "choice_group",
        }
        header_like_count = sum(1 for h in first_row if h in header_fields)

        if header_like_count >= 2:
            mapping: dict[str, int] = {}
            for idx, header in enumerate(first_row):
                if header == "question":
                    header = "question_text"
                if header in {
                    "unit",
                    "marks",
                    "question_text",
                    "co",
                    "bt_level",
                    "difficulty",
                    "status",
                    "choice_group",
                }:
                    mapping[header] = idx

            for row in table.rows[1:]:
                values = [cell.text.strip() for cell in row.cells]
                row_text = " ".join(v for v in values if v)
                row_diagram_paths: list[str] = []
                for cell in row.cells:
                    row_diagram_paths.extend(extract_cell_diagram_paths(cell, doc, image_cache))
                packed_paths = pack_diagram_paths(row_diagram_paths)
                detected = detect_marks_from_heading(row_text)
                if detected is not None:
                    current_table_marks = detected
                    continue

                row_map = build_row_from_mapped_values(mapping, values)
                if "marks" not in row_map or not row_map.get("marks", "").strip():
                    row_map["marks"] = str(current_table_marks) if current_table_marks is not None else ""
                row_map["unit"] = extract_unit_number(row_map.get("unit", ""))
                row_map["question_text"] = normalize_question_text(row_map.get("question_text", ""))
                row_map["bt_level"] = normalize_bt_level(row_map.get("bt_level", ""))
                row_map["diagram_path"] = packed_paths
                if not row_map.get("question_text") or is_non_question_text(row_map["question_text"]):
                    if packed_paths and rows:
                        merged = pack_diagram_paths(split_diagram_paths(rows[-1].get("diagram_path")) + split_diagram_paths(packed_paths))
                        rows[-1]["diagram_path"] = merged
                    continue
                if row_map:
                    rows.append(row_map)
        else:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                row_text = " ".join(v for v in values if v)
                row_diagram_paths: list[str] = []
                for cell in row.cells:
                    row_diagram_paths.extend(extract_cell_diagram_paths(cell, doc, image_cache))
                packed_paths = pack_diagram_paths(row_diagram_paths)
                detected = detect_marks_from_heading(row_text)
                if detected is not None:
                    current_table_marks = detected
                    continue

                parsed = parse_docx_table_row(values, fallback_marks=current_table_marks)
                if parsed:
                    parsed["diagram_path"] = packed_paths
                    rows.append(parsed)
                elif packed_paths and rows:
                    merged = pack_diagram_paths(split_diagram_paths(rows[-1].get("diagram_path")) + split_diagram_paths(packed_paths))
                    rows[-1]["diagram_path"] = merged

    paragraph_lines = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    if paragraph_lines:
        rows.extend(parse_rows_from_lines_state_machine(paragraph_lines))

    return deduplicate_parsed_rows(rows), None


def parse_rows_from_pdf(file_bytes: bytes) -> tuple[list[dict[str, str]], str | None]:
    if PdfReader is None:
        return [], "PDF upload requires pypdf. Install with: pip install pypdf"

    reader = PdfReader(BytesIO(file_bytes))
    rows: list[dict[str, str]] = []
    current_marks: int | None = None

    for page in reader.pages:
        text = page.extract_text() or ""
        for raw_line in text.splitlines():
            line = raw_line.strip()
            if not line:
                continue

            detected = detect_marks_from_heading(line)
            if detected is not None:
                current_marks = detected
                continue

            parsed = parse_structured_line(line)
            if parsed is None:
                parsed = parse_free_text_line(line, fallback_marks=current_marks)
            if parsed is None and current_marks is not None:
                # Common PDF extraction pattern: columns separated by multiple spaces.
                parts = [p.strip() for p in re.split(r"\s{2,}", line) if p.strip()]
                if len(parts) >= 1:
                    parsed = parse_docx_table_row(parts, fallback_marks=current_marks)
            if parsed:
                rows.append(parsed)

    return rows, None


def get_db() -> sqlite3.Connection:
    if "db" not in g:
        g.db = sqlite3.connect(DB_PATH)
        g.db.row_factory = sqlite3.Row
        ensure_schema_extensions(g.db)
    return g.db


def get_user_by_email(email: str) -> sqlite3.Row | None:
    db = get_db()
    return db.execute(
        "SELECT id, username, full_name, email, role, password_hash, is_active FROM users WHERE email = ?",
        (email,),
    ).fetchone()


def get_user_by_id(user_id: int) -> sqlite3.Row | None:
    db = get_db()
    return db.execute(
        "SELECT id, username, full_name, email, role, password_hash, is_active FROM users WHERE id = ?",
        (user_id,),
    ).fetchone()


def is_authenticated() -> bool:
    return bool(session.get("logged_in"))


def is_admin() -> bool:
    return is_authenticated() and session.get("user_role") == "admin"


def require_admin(view_func):
    @wraps(view_func)
    def wrapped(*args, **kwargs):
        if not is_authenticated():
            return redirect(url_for("login"))
        if not is_admin():
            flash("Admin access required.", "error")
            return redirect(url_for("index"))
        return view_func(*args, **kwargs)

    return wrapped


def validate_password_strength(password: str) -> str | None:
    if len(password) < 10:
        return "Password must be at least 10 characters."
    if not re.search(r"[A-Z]", password):
        return "Password must include at least one uppercase letter."
    if not re.search(r"[a-z]", password):
        return "Password must include at least one lowercase letter."
    if not re.search(r"\d", password):
        return "Password must include at least one number."
    if not re.search(r"[^A-Za-z0-9]", password):
        return "Password must include at least one special character."
    return None


@app.before_request
def require_login() -> Any:
    endpoint = request.endpoint or ""
    public_endpoints = {"login", "register", "logout", "static"}

    if endpoint in public_endpoints or endpoint.startswith("static"):
        return None

    if not is_authenticated():
        return redirect(url_for("login"))
    return None


@app.teardown_appcontext
def close_db(_: Any) -> None:
    db = g.pop("db", None)
    if db is not None:
        db.close()


@app.route("/login", methods=["GET", "POST"])
def login() -> str | Any:
    if is_authenticated():
        return redirect(url_for("index"))

    if request.method == "GET":
        return render_template("login.html")

    email = request.form.get("email", "").strip().lower()
    password = request.form.get("password", "")

    if not email or not password:
        flash("Email and password are required.", "error")
        return render_template("login.html")

    user = get_user_by_email(email)
    if not user:
        flash("Invalid email or password.", "error")
        return render_template("login.html")

    if not user["is_active"]:
        flash("Your account is inactive. Contact admin.", "error")
        return render_template("login.html")

    if not check_password_hash(user["password_hash"], password):
        flash("Invalid email or password.", "error")
        return render_template("login.html")

    session["logged_in"] = True
    session["user_id"] = user["id"]
    session["user_email"] = user["email"]
    session["user_name"] = user["full_name"]
    session["user_role"] = user["role"]
    flash(f"Welcome, {user['full_name']}!", "success")
    return redirect(url_for("index"))


@app.route("/register", methods=["GET", "POST"])
def register() -> str | Any:
    if is_authenticated():
        return redirect(url_for("index"))

    if request.method == "GET":
        return render_template("register.html")

    full_name = request.form.get("full_name", "").strip()
    username = request.form.get("username", "").strip()
    email = request.form.get("email", "").strip().lower()
    password = request.form.get("password", "")
    confirm_password = request.form.get("confirm_password", "")

    if not full_name or not username or not email or not password or not confirm_password:
        flash("All fields are required.", "error")
        return render_template("register.html")

    if password != confirm_password:
        flash("Password and confirm password do not match.", "error")
        return render_template("register.html")

    password_error = validate_password_strength(password)
    if password_error:
        flash(password_error, "error")
        return render_template("register.html")

    db = get_db()
    existing = db.execute(
        "SELECT id FROM users WHERE email = ? OR username = ?",
        (email, username),
    ).fetchone()
    if existing:
        flash("Email or username already exists.", "error")
        return render_template("register.html")

    db.execute(
        """
        INSERT INTO users (username, password_hash, full_name, email, role, is_active)
        VALUES (?, ?, ?, ?, 'teacher', 1)
        """,
        (username, generate_password_hash(password), full_name, email),
    )
    db.commit()

    flash("Registration successful. Please login with your email and password.", "success")
    return redirect(url_for("login"))


@app.route("/logout")
def logout() -> Any:
    session.clear()
    flash("Logged out successfully.", "success")
    return redirect(url_for("login"))


@app.route("/admin/users", methods=["GET"])
@require_admin
def admin_users() -> str:
    db = get_db()
    users = db.execute(
        """
        SELECT id, username, full_name, email, role, is_active, created_at, updated_at
        FROM users
        ORDER BY CASE WHEN role = 'admin' THEN 0 ELSE 1 END, id ASC
        """
    ).fetchall()
    return render_template("admin_users.html", users=users)


@app.route("/admin/users/<int:user_id>/reset-password", methods=["POST"])
@require_admin
def admin_reset_password(user_id: int) -> Any:
    target_user = get_user_by_id(user_id)
    if not target_user:
        flash("User not found.", "error")
        return redirect(url_for("admin_users"))

    new_password = request.form.get("new_password", "")
    confirm_password = request.form.get("confirm_password", "")

    if not new_password or not confirm_password:
        flash("New password and confirm password are required.", "error")
        return redirect(url_for("admin_users"))

    if new_password != confirm_password:
        flash("Password and confirm password do not match.", "error")
        return redirect(url_for("admin_users"))

    password_error = validate_password_strength(new_password)
    if password_error:
        flash(password_error, "error")
        return redirect(url_for("admin_users"))

    db = get_db()
    db.execute(
        """
        UPDATE users
        SET password_hash = ?, updated_at = CURRENT_TIMESTAMP
        WHERE id = ?
        """,
        (generate_password_hash(new_password), user_id),
    )
    db.commit()
    flash(f"Password reset successful for {target_user['email']}.", "success")
    return redirect(url_for("admin_users"))


@app.route("/admin/users/<int:user_id>/set-active", methods=["POST"])
@require_admin
def admin_set_user_active(user_id: int) -> Any:
    target_user = get_user_by_id(user_id)
    if not target_user:
        flash("User not found.", "error")
        return redirect(url_for("admin_users"))

    is_active_raw = request.form.get("is_active", "1")
    is_active = 1 if is_active_raw == "1" else 0

    current_user_id = session.get("user_id")
    if current_user_id == user_id and is_active == 0:
        flash("You cannot deactivate your own account.", "error")
        return redirect(url_for("admin_users"))

    db = get_db()
    db.execute(
        """
        UPDATE users
        SET is_active = ?, updated_at = CURRENT_TIMESTAMP
        WHERE id = ?
        """,
        (is_active, user_id),
    )
    db.commit()

    if is_active:
        flash(f"User {target_user['email']} activated.", "success")
    else:
        flash(f"User {target_user['email']} deactivated.", "success")

    return redirect(url_for("admin_users"))


def get_active_question_bank():
    """Get the active question bank from session, fallback to latest available."""
    db = get_db()
    active_bank_id = session.get("active_bank_id")

    if active_bank_id is not None:
        bank = db.execute(
            """
            SELECT qb.*, s.subject_code, s.subject_name, sem.semester_number, sem.academic_year
            FROM question_banks qb
            JOIN subjects s ON qb.subject_id = s.id
            JOIN semesters sem ON qb.semester_id = sem.id
            WHERE qb.id = ?
            """,
            (active_bank_id,),
        ).fetchone()
        if bank:
            return bank
        session.pop("active_bank_id", None)

    bank = db.execute("""
        SELECT qb.*, s.subject_code, s.subject_name, sem.semester_number, sem.academic_year
        FROM question_banks qb
        JOIN subjects s ON qb.subject_id = s.id
        JOIN semesters sem ON qb.semester_id = sem.id
        ORDER BY qb.id DESC
        LIMIT 1
    """).fetchone()
    if bank:
        session["active_bank_id"] = bank["id"]
    return bank


def list_question_banks() -> list[sqlite3.Row]:
    db = get_db()
    return db.execute(
        """
        SELECT qb.id, s.subject_code, s.subject_name, sem.semester_number, sem.academic_year
        FROM question_banks qb
        JOIN subjects s ON qb.subject_id = s.id
        JOIN semesters sem ON qb.semester_id = sem.id
        ORDER BY s.subject_code ASC, sem.semester_number ASC, sem.academic_year DESC
        """
    ).fetchall()


def ensure_default_units(question_bank_id: int) -> None:
    db = get_db()
    existing = db.execute(
        "SELECT COUNT(*) FROM units WHERE question_bank_id = ?",
        (question_bank_id,),
    ).fetchone()[0]
    if existing > 0:
        return

    for unit_number in range(1, 6):
        db.execute(
            """
            INSERT INTO units (question_bank_id, unit_number, unit_name, description)
            VALUES (?, ?, ?, ?)
            """,
            (question_bank_id, unit_number, f"Unit {unit_number}", ""),
        )
    db.commit()


@app.route("/subjects/select", methods=["POST"])
def select_subject_bank() -> Any:
    bank_id_raw = request.form.get("bank_id", "").strip()
    if not bank_id_raw.isdigit():
        flash("Please choose a valid subject.", "error")
        return redirect(url_for("index"))

    bank_id = int(bank_id_raw)
    banks = list_question_banks()
    valid_ids = {b["id"] for b in banks}
    if bank_id not in valid_ids:
        flash("Selected subject is not available.", "error")
        return redirect(url_for("index"))

    session["active_bank_id"] = bank_id
    flash("Active subject changed successfully.", "success")
    return redirect(request.referrer or url_for("index"))


@app.route("/subjects/create", methods=["POST"])
def create_subject_bank() -> Any:
    db = get_db()

    subject_code = request.form.get("subject_code", "").strip().upper()
    subject_name = request.form.get("subject_name", "").strip()
    department = request.form.get("department", "").strip()
    semester_raw = request.form.get("semester_number", "").strip()
    academic_year = request.form.get("academic_year", "").strip()

    if not subject_code or not subject_name or not semester_raw or not academic_year:
        flash("Subject code, subject name, semester, and academic year are required.", "error")
        return redirect(url_for("index"))

    if not semester_raw.isdigit() or int(semester_raw) not in {1, 2, 3, 4, 5, 6, 7, 8}:
        flash("Semester must be between 1 and 8.", "error")
        return redirect(url_for("index"))
    semester_number = int(semester_raw)

    subject = db.execute(
        "SELECT id FROM subjects WHERE subject_code = ?",
        (subject_code,),
    ).fetchone()
    if subject:
        subject_id = subject["id"]
    else:
        db.execute(
            """
            INSERT INTO subjects (subject_code, subject_name, department, total_marks, created_by, is_active)
            VALUES (?, ?, ?, 100, ?, 1)
            """,
            (subject_code, subject_name, department, 1),
        )
        subject_id = db.execute("SELECT last_insert_rowid()").fetchone()[0]

    semester = db.execute(
        "SELECT id FROM semesters WHERE semester_number = ? AND academic_year = ?",
        (semester_number, academic_year),
    ).fetchone()
    if semester:
        semester_id = semester["id"]
    else:
        db.execute(
            """
            INSERT INTO semesters (semester_number, academic_year, is_active)
            VALUES (?, ?, 1)
            """,
            (semester_number, academic_year),
        )
        semester_id = db.execute("SELECT last_insert_rowid()").fetchone()[0]

    bank = db.execute(
        "SELECT id FROM question_banks WHERE subject_id = ? AND semester_id = ?",
        (subject_id, semester_id),
    ).fetchone()
    if bank:
        bank_id = bank["id"]
    else:
        db.execute(
            """
            INSERT INTO question_banks (subject_id, semester_id, created_by, total_questions)
            VALUES (?, ?, ?, 0)
            """,
            (subject_id, semester_id, 1),
        )
        bank_id = db.execute("SELECT last_insert_rowid()").fetchone()[0]

    db.commit()
    ensure_default_units(bank_id)
    session["active_bank_id"] = bank_id
    flash("Subject and bank created/selected successfully.", "success")
    return redirect(url_for("index"))


@app.route("/subjects/delete", methods=["POST"])
def delete_subject_bank() -> Any:
    db = get_db()

    bank_id_raw = request.form.get("bank_id", "").strip()
    if not bank_id_raw.isdigit():
        flash("Please choose a valid subject to delete.", "error")
        return redirect(request.referrer or url_for("index"))

    bank_id = int(bank_id_raw)

    bank = db.execute(
        """
        SELECT qb.id, qb.subject_id, qb.semester_id, s.subject_code, s.subject_name, sem.semester_number, sem.academic_year
        FROM question_banks qb
        JOIN subjects s ON qb.subject_id = s.id
        JOIN semesters sem ON qb.semester_id = sem.id
        WHERE qb.id = ?
        """,
        (bank_id,),
    ).fetchone()
    if not bank:
        flash("Selected subject bank is not available.", "error")
        return redirect(request.referrer or url_for("index"))

    diagram_rows = db.execute(
        "SELECT diagram_path FROM questions WHERE question_bank_id = ?",
        (bank_id,),
    ).fetchall()
    for row in diagram_rows:
        delete_diagram_paths(row["diagram_path"])

    db.execute("DELETE FROM questions WHERE question_bank_id = ?", (bank_id,))
    db.execute("DELETE FROM units WHERE question_bank_id = ?", (bank_id,))
    db.execute("DELETE FROM question_banks WHERE id = ?", (bank_id,))

    remaining_for_subject = db.execute(
        "SELECT COUNT(*) FROM question_banks WHERE subject_id = ?",
        (bank["subject_id"],),
    ).fetchone()[0]
    if remaining_for_subject == 0:
        db.execute("DELETE FROM subjects WHERE id = ?", (bank["subject_id"],))

    remaining_for_semester = db.execute(
        "SELECT COUNT(*) FROM question_banks WHERE semester_id = ?",
        (bank["semester_id"],),
    ).fetchone()[0]
    if remaining_for_semester == 0:
        db.execute("DELETE FROM semesters WHERE id = ?", (bank["semester_id"],))

    db.commit()

    current_active_bank_id = session.get("active_bank_id")
    if current_active_bank_id == bank_id:
        next_bank = db.execute(
            "SELECT id FROM question_banks ORDER BY id ASC LIMIT 1"
        ).fetchone()
        if next_bank:
            session["active_bank_id"] = next_bank["id"]
        else:
            session.pop("active_bank_id", None)

    flash(
        f"Deleted subject bank: {bank['subject_code']} - {bank['subject_name']} (Sem {bank['semester_number']}, {bank['academic_year']}).",
        "success",
    )
    return redirect(request.referrer or url_for("index"))


def fetch_questions_by_marks(marks: int) -> list[sqlite3.Row]:
    """Get approved questions by marks from active bank"""
    db = get_db()
    bank = get_active_question_bank()
    if not bank:
        return []
    
    cur = db.execute(
        """
        SELECT q.id, q.question_text, q.marks, q.co_mapping, q.bt_level, q.difficulty,
               q.unit_id, u.unit_number, q.choice_group, q.diagram_path
        FROM questions q
        LEFT JOIN units u ON q.unit_id = u.id
        WHERE q.question_bank_id = ? AND q.marks = ? AND q.status = 'approved'
        ORDER BY q.id ASC
        """,
        (bank['id'], marks),
    )
    return list(cur.fetchall())


def deduplicate_question_pool(pool: list[sqlite3.Row]) -> list[sqlite3.Row]:
    """Remove duplicate questions by normalized text within a pool."""
    unique_rows: list[sqlite3.Row] = []
    seen_keys: set[str] = set()

    for q in pool:
        normalized = normalize_question_text((q["question_text"] or "")).lower()
        if not normalized:
            continue
        if normalized in seen_keys:
            continue
        seen_keys.add(normalized)
        unique_rows.append(q)

    return unique_rows


def auto_pick_with_teacher_priority(
    all_questions: list[sqlite3.Row],
    selected_ids: set[int],
    required_count: int,
) -> list[sqlite3.Row]:
    """Auto-fill questions from teacher selection or from pool"""
    selected = [q for q in all_questions if q["id"] in selected_ids]
    selected_ids_local = {q["id"] for q in selected}

    if len(selected) < required_count:
        for q in all_questions:
            if q["id"] in selected_ids_local:
                continue
            selected.append(q)
            selected_ids_local.add(q["id"])
            if len(selected) == required_count:
                break

    return selected[:required_count]


@app.route("/")
def index() -> str:
    db = get_db()
    banks = list_question_banks()
    
    bank = get_active_question_bank()
    
    if not bank:
        return render_template(
            "index.html",
            counts={2: 0, 13: 0, 15: 0},
            total=0,
            no_bank=True,
            banks=banks,
            active_bank_id=session.get("active_bank_id"),
        )
    
    counts = {
        2: db.execute(
            "SELECT COUNT(*) FROM questions WHERE question_bank_id = ? AND marks = 2 AND status = 'approved'",
            (bank['id'],)
        ).fetchone()[0],
        13: db.execute(
            "SELECT COUNT(*) FROM questions WHERE question_bank_id = ? AND marks = 13 AND status = 'approved'",
            (bank['id'],)
        ).fetchone()[0],
        15: db.execute(
            "SELECT COUNT(*) FROM questions WHERE question_bank_id = ? AND marks IN (15, 16) AND status = 'approved'",
            (bank['id'],)
        ).fetchone()[0],
    }
    total = db.execute(
        "SELECT COUNT(*) FROM questions WHERE question_bank_id = ?",
        (bank['id'],)
    ).fetchone()[0]
    
    return render_template(
        "index.html",
        counts=counts,
        total=total,
        bank=bank,
        no_bank=False,
        banks=banks,
        active_bank_id=bank["id"],
    )


@app.route("/questions", methods=["GET", "POST"])
def questions() -> str | Any:
    db = get_db()
    
    bank = get_active_question_bank()
    if not bank:
        flash("No question bank found. Please initialize database first.", "error")
        return redirect(url_for("index"))

    if request.method == "POST":
        unit_number = request.form.get("unit", "1").strip()
        marks = int(request.form.get("marks", "0"))
        question_text = request.form.get("question_text", "").strip()
        co = request.form.get("co", "").strip()
        bt_level = request.form.get("bt_level", "").strip()
        difficulty = request.form.get("difficulty", "easy").strip()
        status = request.form.get("status", "approved").strip()
        choice_group = request.form.get("choice_group", "").strip()
        diagram_file = request.files.get("question_diagram")
        diagram_path = save_diagram_file(diagram_file)

        if diagram_file and diagram_file.filename and diagram_path is None:
            flash("Diagram must be an image file: png, jpg, jpeg, webp, svg.", "error")
            return redirect(url_for("questions"))

        if marks not in {2, 13, 15, 16}:
            flash("Marks must be 2, 13, 15, or 16.", "error")
            return redirect(url_for("questions"))

        if not question_text:
            flash("Question text is required.", "error")
            return redirect(url_for("questions"))

        # Get or create unit
        unit = db.execute(
            "SELECT id FROM units WHERE question_bank_id = ? AND unit_number = ?",
            (bank['id'], int(unit_number))
        ).fetchone()
        
        if not unit:
            flash(f"Unit {unit_number} not found in question bank.", "error")
            return redirect(url_for("questions"))

        db.execute(
            """
            INSERT INTO questions 
            (question_bank_id, unit_id, marks, question_text, co_mapping, bt_level, difficulty, status, choice_group, created_by, diagram_path)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                bank['id'],
                unit['id'],
                marks,
                question_text,
                co,
                bt_level,
                difficulty,
                status,
                choice_group,
                1,
                diagram_path,
            ),
        )
        db.commit()
        flash("Question added successfully.", "success")
        return redirect(url_for("questions"))

    mark_filter = request.args.get("marks", "all")

    if mark_filter in {"2", "13", "15", "16"}:
        rows = db.execute("""
            SELECT q.*, u.unit_number
            FROM questions q
            LEFT JOIN units u ON q.unit_id = u.id
            WHERE q.question_bank_id = ? AND q.marks = ? 
            ORDER BY q.id DESC
        """, (bank['id'], int(mark_filter))).fetchall()
    else:
        rows = db.execute("""
            SELECT q.*, u.unit_number
            FROM questions q
            LEFT JOIN units u ON q.unit_id = u.id
            WHERE q.question_bank_id = ?
            ORDER BY q.id DESC
        """, (bank['id'],)).fetchall()

    # Group rows by marks for mark-wise rendering in bank.html
    rows_by_marks = {2: [], 13: [], 15: [], 16: []}
    for q in rows:
        marks_value = q["marks"] if "marks" in q.keys() else None
        if marks_value in rows_by_marks:
            rows_by_marks[marks_value].append(q)

    # Get available units
    units = db.execute(
        "SELECT unit_number FROM units WHERE question_bank_id = ? ORDER BY unit_number",
        (bank['id'],)
    ).fetchall()
    banks = list_question_banks()

    return render_template(
        "bank.html",
        rows=rows,
        rows_by_marks=rows_by_marks,
        mark_filter=mark_filter,
        bank=bank,
        units=units,
        banks=banks,
        active_bank_id=bank["id"],
    )


@app.route("/questions/<int:question_id>/delete", methods=["POST"])
def delete_question(question_id: int) -> Any:
    db = get_db()
    bank = get_active_question_bank()
    if not bank:
        flash("No active question bank found.", "error")
        return redirect(url_for("index"))

    row = db.execute(
        "SELECT diagram_path FROM questions WHERE id = ? AND question_bank_id = ?",
        (question_id, bank["id"]),
    ).fetchone()
    if not row:
        flash("Question not found in active subject.", "error")
        return redirect(url_for("questions"))

    db.execute(
        "DELETE FROM questions WHERE id = ? AND question_bank_id = ?",
        (question_id, bank["id"]),
    )
    db.commit()
    if row:
        delete_diagram_paths(row["diagram_path"])
    flash("Question deleted.", "success")
    return redirect(url_for("questions"))


@app.route("/questions/delete-bulk", methods=["POST"])
def delete_questions_bulk() -> Any:
    db = get_db()
    bank = get_active_question_bank()
    if not bank:
        flash("No active question bank found.", "error")
        return redirect(url_for("index"))

    raw_ids = request.form.getlist("question_ids")
    ids: list[int] = []
    for raw in raw_ids:
        if raw.isdigit():
            ids.append(int(raw))

    if not ids:
        flash("Select at least one question to delete.", "error")
        return redirect(url_for("questions"))

    placeholders = ",".join("?" for _ in ids)
    params: tuple[Any, ...] = (bank["id"], *ids)

    rows = db.execute(
        f"""
        SELECT id, diagram_path
        FROM questions
        WHERE question_bank_id = ? AND id IN ({placeholders})
        """,
        params,
    ).fetchall()

    if not rows:
        flash("No matching questions found in active subject.", "error")
        return redirect(url_for("questions"))

    db.execute(
        f"DELETE FROM questions WHERE question_bank_id = ? AND id IN ({placeholders})",
        params,
    )
    db.commit()

    for row in rows:
        delete_diagram_paths(row["diagram_path"])

    flash(f"Deleted {len(rows)} questions.", "success")
    return redirect(url_for("questions"))


@app.route("/questions/<int:question_id>/edit", methods=["GET", "POST"])
def edit_question(question_id: int) -> str | Any:
    db = get_db()
    bank = get_active_question_bank()
    if not bank:
        flash("No question bank found. Please initialize database first.", "error")
        return redirect(url_for("index"))

    question = db.execute(
        """
        SELECT q.*, u.unit_number
        FROM questions q
        LEFT JOIN units u ON q.unit_id = u.id
        WHERE q.id = ? AND q.question_bank_id = ?
        """,
        (question_id, bank["id"]),
    ).fetchone()

    if not question:
        flash("Question not found in active bank.", "error")
        return redirect(url_for("questions"))

    units = db.execute(
        "SELECT unit_number FROM units WHERE question_bank_id = ? ORDER BY unit_number",
        (bank["id"],),
    ).fetchall()

    current_diagrams = split_diagram_paths(question["diagram_path"])

    if request.method == "GET":
        return render_template(
            "question_edit.html",
            question=question,
            units=units,
            bank=bank,
            current_diagrams=current_diagrams,
        )

    unit_number = request.form.get("unit", "1").strip()
    marks_raw = request.form.get("marks", "0").strip()
    question_text = request.form.get("question_text", "").strip()
    co = request.form.get("co", "").strip()
    bt_level = request.form.get("bt_level", "").strip()
    difficulty = request.form.get("difficulty", "medium").strip().lower()
    status = request.form.get("status", "approved").strip().lower()
    choice_group = request.form.get("choice_group", "").strip()
    remove_diagram = request.form.get("remove_diagram") == "on"
    diagram_file = request.files.get("question_diagram")

    try:
        marks = int(marks_raw)
    except ValueError:
        flash("Marks must be a valid number.", "error")
        return redirect(url_for("edit_question", question_id=question_id))

    if marks not in {2, 13, 15, 16}:
        flash("Marks must be 2, 13, 15, or 16.", "error")
        return redirect(url_for("edit_question", question_id=question_id))

    if not question_text:
        flash("Question text is required.", "error")
        return redirect(url_for("edit_question", question_id=question_id))

    if difficulty not in {"easy", "medium", "hard"}:
        flash("Difficulty must be easy, medium, or hard.", "error")
        return redirect(url_for("edit_question", question_id=question_id))

    if status not in {"approved", "draft"}:
        flash("Status must be approved or draft.", "error")
        return redirect(url_for("edit_question", question_id=question_id))

    unit = db.execute(
        "SELECT id FROM units WHERE question_bank_id = ? AND unit_number = ?",
        (bank["id"], int(unit_number)),
    ).fetchone()
    if not unit:
        flash(f"Unit {unit_number} not found in active bank.", "error")
        return redirect(url_for("edit_question", question_id=question_id))

    new_diagram_path = question["diagram_path"]

    if remove_diagram and new_diagram_path:
        delete_diagram_paths(new_diagram_path)
        new_diagram_path = None

    if diagram_file and diagram_file.filename:
        uploaded_path = save_diagram_file(diagram_file)
        if uploaded_path is None:
            flash("Diagram must be an image file: png, jpg, jpeg, webp, svg.", "error")
            return redirect(url_for("edit_question", question_id=question_id))
        if new_diagram_path and new_diagram_path != uploaded_path:
            delete_diagram_paths(new_diagram_path)
        new_diagram_path = uploaded_path

    db.execute(
        """
        UPDATE questions
        SET unit_id = ?, marks = ?, question_text = ?, co_mapping = ?, bt_level = ?,
            difficulty = ?, status = ?, choice_group = ?, diagram_path = ?
        WHERE id = ? AND question_bank_id = ?
        """,
        (
            unit["id"],
            marks,
            question_text,
            co,
            bt_level,
            difficulty,
            status,
            choice_group,
            new_diagram_path,
            question_id,
            bank["id"],
        ),
    )
    db.commit()

    flash("Question updated successfully.", "success")
    return redirect(url_for("questions"))


@app.route("/questions/upload", methods=["POST"])
def upload_questions() -> Any:
    bank = get_active_question_bank()
    if not bank:
        flash("No question bank found. Please initialize database first.", "error")
        return redirect(url_for("index"))

    file = request.files.get("questions_file")
    if not file or not file.filename:
        flash("Please choose a file to upload.", "error")
        return redirect(url_for("questions"))

    extension = Path(file.filename).suffix.lower()
    if extension not in {".csv", ".docx", ".pdf"}:
        flash("Supported formats: .csv, .docx, .pdf", "error")
        return redirect(url_for("questions"))

    file_bytes = file.read()
    rows_to_import: list[dict[str, str]] = []

    if extension == ".csv":
        try:
            csv_text = file_bytes.decode("utf-8-sig")
        except UnicodeDecodeError:
            flash("CSV must be UTF-8 encoded.", "error")
            return redirect(url_for("questions"))

        reader = csv.DictReader(csv_text.splitlines())
        required_columns = {"marks", "question_text"}

        if not reader.fieldnames:
            flash("CSV appears empty or invalid.", "error")
            return redirect(url_for("questions"))

        normalized_headers = {normalize_column_name(h) for h in reader.fieldnames}
        missing_columns = required_columns - normalized_headers
        if missing_columns:
            flash(
                f"Missing required columns: {', '.join(sorted(missing_columns))}",
                "error",
            )
            return redirect(url_for("questions"))

        for row in reader:
            normalized_row: dict[str, str] = {}
            for key, value in row.items():
                normalized_row[normalize_column_name(key)] = (value or "").strip()
            if "question" in normalized_row and "question_text" not in normalized_row:
                normalized_row["question_text"] = normalized_row["question"]
            rows_to_import.append(normalized_row)

    elif extension == ".docx":
        rows_to_import, parse_error = parse_rows_from_docx(file_bytes)
        if parse_error:
            flash(parse_error, "error")
            return redirect(url_for("questions"))
    else:
        rows_to_import, parse_error = parse_rows_from_pdf(file_bytes)
        if parse_error:
            flash(parse_error, "error")
            return redirect(url_for("questions"))

    staged_rows: list[dict[str, str]] = []
    for row in rows_to_import:
        inferred_marks = infer_total_marks_from_text(row.get("question_text") or "", None)
        normalized = {
            "unit": extract_unit_number((row.get("unit") or "").strip()) or "1",
            "marks": str(extract_marks((row.get("marks") or "").strip()) or inferred_marks or ""),
            "question_text": normalize_question_text_preserve_structure(row.get("question_text") or ""),
            "co": (row.get("co") or "").strip(),
            "bt_level": normalize_bt_level((row.get("bt_level") or "").strip()),
            "difficulty": (row.get("difficulty") or "medium").strip().lower() or "medium",
            "status": (row.get("status") or "approved").strip().lower() or "approved",
            "choice_group": (row.get("choice_group") or "").strip(),
            "diagram_path": (row.get("diagram_path") or "").strip(),
        }
        if not normalized["marks"]:
            normalized["marks"] = "2"
        if not normalized["unit"]:
            normalized["unit"] = "1"
        if normalized["question_text"] and not is_non_question_text(normalized["question_text"]):
            staged_rows.append(normalized)

    staged_rows = deduplicate_parsed_rows(staged_rows)

    if not staged_rows:
        flash("No importable rows found in the uploaded file.", "error")
        return redirect(url_for("questions"))

    old_token = session.get("pending_upload_token")
    delete_staged_upload(old_token)

    token = save_staged_upload(staged_rows)
    session["pending_upload_token"] = token

    flash(f"Parsed {len(staged_rows)} rows. Review and correct before saving.", "success")
    return redirect(url_for("preview_upload_questions"))


@app.route("/questions/upload/preview", methods=["GET"])
def preview_upload_questions() -> str | Any:
    bank = get_active_question_bank()
    if not bank:
        flash("No active question bank found.", "error")
        return redirect(url_for("index"))

    token = session.get("pending_upload_token")
    rows = load_staged_upload(token) if token else None
    if not rows:
        flash("No parsed upload data found. Please upload a file first.", "error")
        return redirect(url_for("questions"))

    db = get_db()
    units = db.execute(
        "SELECT unit_number FROM units WHERE question_bank_id = ? ORDER BY unit_number",
        (bank["id"],),
    ).fetchall()

    return render_template(
        "upload_preview.html",
        rows=rows,
        units=units,
        bank=bank,
        banks=list_question_banks(),
        active_bank_id=bank["id"],
    )


@app.route("/questions/upload/cancel", methods=["POST"])
def cancel_upload_questions() -> Any:
    token = session.pop("pending_upload_token", None)
    delete_staged_upload(token, cleanup_diagrams=True)
    flash("Upload preview canceled.", "success")
    return redirect(url_for("questions"))


@app.route("/questions/upload/confirm", methods=["POST"])
def confirm_upload_questions() -> Any:
    db = get_db()
    bank = get_active_question_bank()
    if not bank:
        flash("No active question bank found.", "error")
        return redirect(url_for("index"))

    token = session.get("pending_upload_token")
    rows = load_staged_upload(token) if token else None
    if not rows:
        flash("No parsed upload data found. Please upload again.", "error")
        return redirect(url_for("questions"))

    units = db.execute(
        "SELECT id, unit_number FROM units WHERE question_bank_id = ?",
        (bank["id"],),
    ).fetchall()
    unit_map = {str(u["unit_number"]): u["id"] for u in units}

    try:
        total_rows = int(request.form.get("total_rows", "0"))
    except ValueError:
        total_rows = 0

    inserted_count = 0
    skipped_count = 0
    error_messages: list[str] = []

    for idx in range(total_rows):
        if request.form.get(f"include_{idx}") != "on":
            skipped_count += 1
            continue

        unit_raw = extract_unit_number(request.form.get(f"unit_{idx}", "")) or request.form.get(f"unit_{idx}", "1").strip() or "1"
        marks_raw = request.form.get(f"marks_{idx}", "").strip()
        question_text = normalize_question_text_preserve_structure(request.form.get(f"question_text_{idx}", ""))
        co_mapping = request.form.get(f"co_{idx}", "").strip()
        bt_level = normalize_bt_level(request.form.get(f"bt_level_{idx}", "").strip())
        difficulty = request.form.get(f"difficulty_{idx}", "medium").strip().lower()
        status = request.form.get(f"status_{idx}", "approved").strip().lower()
        choice_group = request.form.get(f"choice_group_{idx}", "").strip()
        
        # Determine diagram path: check form first (for split clones), fallback to original staged rows
        diagram_path_form = request.form.get(f"diagram_path_{idx}")
        if diagram_path_form is not None:
            diagram_path = diagram_path_form.strip()
        else:
            source_row = rows[idx] if idx < len(rows) else {}
            diagram_path = (source_row.get("diagram_path") or "").strip() if isinstance(source_row, dict) else ""

        if is_non_question_text(question_text):
            skipped_count += 1
            continue

        if not marks_raw or not question_text:
            skipped_count += 1
            error_messages.append(f"Row {idx + 1}: marks and question text are required.")
            continue

        marks = extract_marks(marks_raw)
        if marks is None or marks not in {2, 13, 15, 16}:
            skipped_count += 1
            error_messages.append(f"Row {idx + 1}: marks must be 2, 13, 15, or 16.")
            continue

        unit_id = unit_map.get(unit_raw)
        if not unit_id:
            skipped_count += 1
            error_messages.append(f"Row {idx + 1}: unit {unit_raw} not found in active subject.")
            continue

        if difficulty not in {"easy", "medium", "hard"}:
            skipped_count += 1
            error_messages.append(f"Row {idx + 1}: difficulty must be easy, medium, or hard.")
            continue

        if status not in {"approved", "draft"}:
            skipped_count += 1
            error_messages.append(f"Row {idx + 1}: status must be approved or draft.")
            continue

        db.execute(
            """
            INSERT INTO questions
            (question_bank_id, unit_id, marks, question_text, co_mapping, bt_level, difficulty, status, choice_group, created_by, diagram_path)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                bank["id"],
                unit_id,
                marks,
                question_text,
                co_mapping,
                bt_level,
                difficulty,
                status,
                choice_group,
                1,
                diagram_path,
            ),
        )
        inserted_count += 1

    if inserted_count > 0:
        db.commit()

    delete_staged_upload(token, cleanup_diagrams=False)
    session.pop("pending_upload_token", None)

    if inserted_count:
        flash(f"Upload saved: {inserted_count} inserted, {skipped_count} skipped.", "success")
    else:
        flash("No questions inserted from preview. Please correct and retry upload.", "error")

    if error_messages:
        flash("Some rows failed: " + " | ".join(error_messages[:3]), "error")

    return redirect(url_for("questions"))


@app.route("/generate", methods=["GET", "POST"])
def generate() -> str | Any:
    bank = get_active_question_bank()
    if not bank:
        flash("No question bank found. Please initialize database.", "error")
        return redirect(url_for("index"))

    def filter_by_lessons(pool: list[sqlite3.Row], lessons: set[int]) -> list[sqlite3.Row]:
        if not lessons:
            return pool
        return [q for q in pool if q["unit_number"] in lessons]

    part_a_pool_all = deduplicate_question_pool(fetch_questions_by_marks(2))
    part_b_pool_all = deduplicate_question_pool(fetch_questions_by_marks(13))
    db = get_db()
    part_c_pool_all = deduplicate_question_pool(list(
        db.execute(
            """
            SELECT q.id, q.question_text, q.marks, q.co_mapping, q.bt_level, q.difficulty,
                   q.unit_id, u.unit_number, q.choice_group, q.diagram_path
            FROM questions q
            LEFT JOIN units u ON q.unit_id = u.id
                 WHERE q.question_bank_id = ? AND q.marks IN (15, 16) AND q.status = 'approved'
            ORDER BY q.id ASC
            """,
            (bank['id'],),
        ).fetchall()
    ))

    units = db.execute(
        "SELECT unit_number FROM units WHERE question_bank_id = ? ORDER BY unit_number",
        (bank["id"],),
    ).fetchall()

    if request.method == "GET":
        selected_lessons: set[int] = set()
        today = ""
        return render_template(
            "generate.html",
            part_a_pool=part_a_pool_all,
            part_b_pool=part_b_pool_all,
            part_c_pool=part_c_pool_all,
            bank=bank,
            units=units,
            selected_lessons=selected_lessons,
            selected_question_ids=set(),
            today=today,
            college_name="",
            department="",
            assessment_type="",
            semester=f"Semester {bank['semester_number']}",
            exam_type="Model Examination",
            exam_date=today,
            exam_time="",
            part_a_count=10,
            banks=list_question_banks(),
            active_bank_id=bank["id"],
        )

    college_name = request.form.get("college_name", "").strip()
    department = request.form.get("department", "").strip()
    semester = request.form.get("semester", "").strip()
    assessment_type = request.form.get("assessment_type", "").strip()
    exam_type = assessment_type or request.form.get("exam_type", "").strip()
    exam_date = request.form.get("exam_date", "").strip()
    exam_time = request.form.get("exam_time", "").strip()

    part_a_count_raw = request.form.get("part_a_count", "10").strip()
    try:
        part_a_count = max(1, min(20, int(part_a_count_raw)))
    except ValueError:
        part_a_count = 10

    selected_lessons = {int(v) for v in request.form.getlist("selected_lessons") if v.isdigit()}
    part_a_pool = filter_by_lessons(part_a_pool_all, selected_lessons)
    part_b_pool = filter_by_lessons(part_b_pool_all, selected_lessons)
    part_c_pool = filter_by_lessons(part_c_pool_all, selected_lessons)

    selected_ids = {int(i) for i in request.form.getlist("selected_question_ids") if i.isdigit()}

    selected_part_a = auto_pick_with_teacher_priority(part_a_pool, selected_ids, part_a_count)
    selected_part_b = auto_pick_with_teacher_priority(part_b_pool, selected_ids, 10)
    selected_part_c = auto_pick_with_teacher_priority(part_c_pool, selected_ids, 2)

    def render_generate_with_state() -> str:
        return render_template(
            "generate.html",
            part_a_pool=part_a_pool,
            part_b_pool=part_b_pool,
            part_c_pool=part_c_pool,
            bank=bank,
            units=units,
            selected_lessons=selected_lessons,
            selected_question_ids=selected_ids,
            today=exam_date,
            college_name=college_name,
            department=department,
            assessment_type=assessment_type,
            semester=semester,
            exam_type=exam_type,
            exam_date=exam_date,
            exam_time=exam_time,
            part_a_count=part_a_count,
            banks=list_question_banks(),
            active_bank_id=bank["id"],
        )

    if not college_name or not department or not exam_date:
        flash("Please fill college name, department, and exam date.", "error")
        return render_generate_with_state()

    if len(selected_part_a) < part_a_count:
        flash(
            f"Not enough 2-mark questions. Need {part_a_count}, found {len(selected_part_a)}.",
            "error",
        )
        return render_generate_with_state()

    if len(selected_part_b) < 10:
        flash("Not enough 13-mark questions. Need 10 unique questions for 5 OR pairs.", "error")
        return render_generate_with_state()

    if len(selected_part_c) < 2:
        flash("Not enough 15/16-mark questions. Need 2 questions for either/or.", "error")
        return render_generate_with_state()

    paper_id = save_generated_paper(
        bank=bank,
        college_name=college_name,
        department=department,
        semester=semester,
        exam_type=exam_type,
        exam_date=exam_date,
        exam_time=exam_time,
        part_a_count=part_a_count,
        selected_part_a=selected_part_a,
        selected_part_b=selected_part_b,
        selected_part_c=selected_part_c,
    )
    flash("Question paper generated and saved in Recent History.", "success")
    return redirect(url_for("view_saved_paper", paper_id=paper_id))


@app.route("/papers/recent", methods=["GET"])
def recent_papers() -> str:
    db = get_db()
    bank = get_active_question_bank()

    if bank:
        rows = db.execute(
            """
            SELECT qp.id, qp.paper_title, qp.exam_type, qp.exam_date, qp.exam_time, qp.created_at,
                   s.subject_code, s.subject_name, sem.semester_number, sem.academic_year
            FROM question_papers qp
            JOIN subjects s ON qp.subject_id = s.id
            JOIN semesters sem ON qp.semester_id = sem.id
            WHERE qp.subject_id = ? AND qp.semester_id = ?
            ORDER BY qp.created_at DESC, qp.id DESC
            LIMIT 50
            """,
            (bank["subject_id"], bank["semester_id"]),
        ).fetchall()
    else:
        rows = db.execute(
            """
            SELECT qp.id, qp.paper_title, qp.exam_type, qp.exam_date, qp.exam_time, qp.created_at,
                   s.subject_code, s.subject_name, sem.semester_number, sem.academic_year
            FROM question_papers qp
            JOIN subjects s ON qp.subject_id = s.id
            JOIN semesters sem ON qp.semester_id = sem.id
            ORDER BY qp.created_at DESC, qp.id DESC
            LIMIT 50
            """
        ).fetchall()

    return render_template(
        "papers_history.html",
        papers=rows,
        bank=bank,
        banks=list_question_banks(),
        active_bank_id=bank["id"] if bank else session.get("active_bank_id"),
    )


@app.route("/papers/<int:paper_id>", methods=["GET"])
def view_saved_paper(paper_id: int) -> str | Any:
    context = get_paper_render_context(paper_id)
    if not context:
        flash("Saved paper not found.", "error")
        return redirect(url_for("recent_papers"))
    return render_template("paper.html", **context)


@app.route("/papers/<int:paper_id>/save-edits", methods=["POST"])
def save_saved_paper_edits(paper_id: int) -> Any:
    payload = request.get_json(silent=True) or {}
    edits = payload.get("edits")
    if not isinstance(edits, list):
        return jsonify({"ok": False, "error": "Invalid edits payload."}), 400

    db = get_db()
    exists = db.execute("SELECT id FROM question_papers WHERE id = ?", (paper_id,)).fetchone()
    if not exists:
        return jsonify({"ok": False, "error": "Paper not found."}), 404

    db.execute(
        "UPDATE question_papers SET edited_payload_json = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
        (json.dumps(edits, ensure_ascii=True), paper_id),
    )
    db.commit()
    return jsonify({"ok": True})


@app.route("/papers/<int:paper_id>/clear-edits", methods=["POST"])
def clear_saved_paper_edits(paper_id: int) -> Any:
    db = get_db()
    exists = db.execute("SELECT id FROM question_papers WHERE id = ?", (paper_id,)).fetchone()
    if not exists:
        return jsonify({"ok": False, "error": "Paper not found."}), 404

    db.execute(
        "UPDATE question_papers SET edited_payload_json = NULL, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
        (paper_id,),
    )
    db.commit()
    return jsonify({"ok": True})


@app.route("/papers/<int:paper_id>/download-word", methods=["GET"])
def download_saved_paper_word(paper_id: int) -> Any:
    context = get_paper_render_context(paper_id)
    if not context:
        flash("Saved paper not found.", "error")
        return redirect(url_for("recent_papers"))

    if Document is None:
        flash("Word export requires python-docx package.", "error")
        return redirect(url_for("view_saved_paper", paper_id=paper_id))

    stream = build_paper_docx_bytes(context)
    filename = f"paper_{paper_id}.docx"
    return send_file(
        stream,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route("/seed-demo", methods=["POST"])
def seed_demo() -> Any:
    db = get_db()
    
    bank = get_active_question_bank()
    if not bank:
        flash("No question bank found.", "error")
        return redirect(url_for("index"))
    
    existing = db.execute(
        "SELECT COUNT(*) FROM questions WHERE question_bank_id = ?",
        (bank['id'],)
    ).fetchone()[0]
    
    if existing > 0:
        flash("Questions already exist in bank. Demo seed skipped.", "error")
        return redirect(url_for("questions"))

    # Get units in this bank
    units = db.execute(
        "SELECT id, unit_number FROM units WHERE question_bank_id = ? ORDER BY unit_number",
        (bank['id'],)
    ).fetchall()
    
    unit_ids = {u['unit_number']: u['id'] for u in units}

    # Add 2-mark questions (distribute across units)
    for i in range(1, 16):
        unit_num = ((i - 1) % 5) + 1
        unit_id = unit_ids.get(unit_num)
        if unit_id:
            db.execute(
                """
                INSERT INTO questions 
                (question_bank_id, unit_id, marks, question_text, co_mapping, bt_level, difficulty, status, created_by)
                VALUES (?, ?, 2, ?, ?, ?, ?, 'approved', 1)
                """,
                (bank['id'], unit_id, f"Define concept {i} with example.", "CO1", "K2", "easy"),
            )

    # Add 13-mark questions
    for i in range(1, 9):
        unit_num = ((i - 1) % 5) + 1
        unit_id = unit_ids.get(unit_num)
        if unit_id:
            db.execute(
                """
                INSERT INTO questions 
                (question_bank_id, unit_id, marks, question_text, co_mapping, bt_level, difficulty, status, created_by)
                VALUES (?, ?, 13, ?, ?, ?, ?, 'approved', 1)
                """,
                (bank['id'], unit_id, f"Explain problem solving method {i} with neat diagram.", "CO3", "L3", "medium"),
            )

    # Add 15-mark questions with choice grouping
    for i in range(1, 5):
        unit_num = ((i - 1) % 5) + 1
        unit_id = unit_ids.get(unit_num)
        if unit_id:
            choice_group = f"C{i}"
            db.execute(
                """
                INSERT INTO questions 
                (question_bank_id, unit_id, marks, question_text, co_mapping, bt_level, difficulty, status, choice_group, created_by)
                VALUES (?, ?, 15, ?, ?, ?, ?, 'approved', ?, 1)
                """,
                (bank['id'], unit_id, f"Discuss advanced topic {i} part (a) in detail.", "CO4", "L4", "hard", choice_group),
            )

    db.commit()
    flash("Demo questions added successfully.", "success")
    return redirect(url_for("questions"))


if __name__ == "__main__":
    # Do NOT call init_db() here - database_schema.py already initialized it
    app.run(debug=True)
