from docx import Document

try:
    doc = Document(r"C:\Users\PRAJEEN\AppData\Local\Packages\5319275A.WhatsAppDesktop_cv1g1gvanyjgm\LocalState\sessions\380D00BC6FE1B5D7E953E69463CF78C28F7FFBA2\transfers\2026-12\CT1 DM.docx")
    
    print("=" * 80)
    print("DOCUMENT ANALYSIS - COLLEGE QUESTION PAPER TEMPLATE")
    print("=" * 80)
    
    print("\n1. DOCUMENT STRUCTURE (First 30 paragraphs):")
    print("-" * 80)
    
    for i, para in enumerate(doc.paragraphs[:30]):
        text = para.text.strip()
        if text:
            alignment = para.alignment  # 0=left, 1=center, 2=right
            bold = False
            italic = False
            font_size = None
            font_name = None
            
            if para.runs:
                bold = para.runs[0].bold
                italic = para.runs[0].italic
                if para.runs[0].font.size:
                    font_size = para.runs[0].font.size.pt
                if para.runs[0].font.name:
                    font_name = para.runs[0].font.name
            
            align_map = {0: "LEFT", 1: "CENTER", 2: "RIGHT"}
            print(f"\n[{i}] {align_map.get(alignment, 'LEFT')} | Font:{font_name} {font_size}pt | Bold:{bold} Italic:{italic}")
            print(f"    {text[:90]}")
    
    print("\n\n2. TABLES FOUND:", len(doc.tables))
    
    print("\n3. PAGE MARGINS AND SETUP:")
    print("-" * 80)
    for section in doc.sections:
        margins = section.margins
        print(f"Top margin: {margins.top.inches:.2f}\" ({margins.top.pt:.0f}pt)")
        print(f"Left margin: {margins.left.inches:.2f}\" ({margins.left.pt:.0f}pt)")
        print(f"Right margin: {margins.right.inches:.2f}\" ({margins.right.pt:.0f}pt)")
        print(f"Bottom margin: {margins.bottom.inches:.2f}\" ({margins.bottom.pt:.0f}pt)")
        print(f"Page width: {section.page_width.inches:.2f}\"")
        print(f"Page height: {section.page_height.inches:.2f}\"")
    
    print("\n\n4. FULL FIRST SECTION TEXT:")
    print("-" * 80)
    for para in doc.paragraphs[:20]:
        if para.text.strip():
            print(para.text)
    
except Exception as e:
    print(f"Error: {e}")
    import traceback
    traceback.print_exc()
